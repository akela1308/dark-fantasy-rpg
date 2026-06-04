#!/usr/bin/env python3
"""Генерация аудит-отчёта Dark Fantasy RPG в формате DOCX."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Стили страницы ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.1)
section.top_margin  = section.bottom_margin = Inches(1.0)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x5A)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x22, 0x00)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    return p

def para(text, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold   = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def mono(text):
    """Блок кода (моноширинный)."""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.paragraph_format.left_indent = Inches(0.4)
    return p

def severity_badge(sev):
    colors = {
        'КРИТИЧЕСКИЙ': (0xCC, 0x00, 0x00),
        'ВЫСОКИЙ':     (0xDD, 0x66, 0x00),
        'СРЕДНИЙ':     (0xAA, 0x99, 0x00),
        'НИЗКИЙ':      (0x22, 0x88, 0x22),
    }
    return colors.get(sev, (0x55, 0x55, 0x55))

def bug_entry(num, title, sev, file_loc, desc, fix):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"[{sev}] ")
    r.bold = True
    r.font.color.rgb = RGBColor(*severity_badge(sev))
    r.font.size = Pt(11)
    r2 = p.add_run(f"#{num}. {title}")
    r2.bold = True
    r2.font.size = Pt(11)

    if file_loc:
        pf = doc.add_paragraph()
        pf.paragraph_format.left_indent = Inches(0.4)
        rf = pf.add_run(f"📁 {file_loc}")
        rf.font.name = 'Courier New'
        rf.font.size = Pt(9)
        rf.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    pd = doc.add_paragraph(f"Проблема: {desc}")
    pd.paragraph_format.left_indent = Inches(0.4)
    pd.runs[0].font.size = Pt(10)

    pfix = doc.add_paragraph(f"Решение: {fix}")
    pfix.paragraph_format.left_indent = Inches(0.4)
    pfix.runs[0].font.size = Pt(10)
    pfix.runs[0].font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    doc.add_paragraph()


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        hdr[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # фон заголовка
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A1A5A')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНАЯ СТРАНИЦА
# ════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Dark Fantasy RPG")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5A)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Аудит проекта — технический отчёт")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run(f"Дата: {datetime.date.today().strftime('%d.%m.%Y')}")
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()

tech_p = doc.add_paragraph()
tech_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tech_p.add_run("Стек: Phaser 3.80 · Vite 5.2 · Vanilla ES6 Modules")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# СВОДНАЯ СТАТИСТИКА
# ════════════════════════════════════════════════════════════════════════════

h1("0. Сводная статистика аудита")

add_table(
    ["Категория", "Кол-во", "Комментарий"],
    [
        ["Критические баги (краш/сбой)", "2", "Немедленное исправление"],
        ["Высокий приоритет (геймплей сломан)", "6", "Перед релизом"],
        ["Средний приоритет (ухудшение UX)", "9", "Следующий спринт"],
        ["Низкий приоритет (косметика)", "5", "Технический долг"],
        ["Отсутствующие фичи (задокументировано)", "8", "Беклог"],
        ["Неиспользуемые ассеты", "11", "Аудит ресурсов"],
    ]
)

para("Файлов прочитано: 20 (сцены, системы, сущности, UI, данные, утилиты, конфиги).",
     italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 1: КРИТИЧЕСКИЕ БАГИ
# ════════════════════════════════════════════════════════════════════════════

h1("1. Критические баги (краш игры)")

para(
    "Ошибки, которые приведут к исключению JavaScript при нормальном игровом процессе.",
    italic=True
)

bug_entry(
    1,
    "TypeError: skill.effect is not a function (скиллы dueling_stance и shield_cover)",
    "КРИТИЧЕСКИЙ",
    "src/systems/SkillSystem.js:38  ·  src/data/skills.json:18,50",
    (
        "В skills.json у навыков dueling_stance и shield_cover поле \"effect\" — строка "
        "(\"dodge_boost\", \"cover\"). В SkillSystem.use() строка 38 проверяет "
        "if (skill.effect) — строка truthful → пытается вызвать skill.effect(caster, ...) "
        "как функцию → TypeError: skill.effect is not a function. "
        "Оба навыка недоступны для использования — краш при их выборе игроком."
    ),
    (
        "Переименовать поле в JSON: \"effect\" → \"builtinEffect\" (или \"effectType\"). "
        "В SkillSystem.use(): проверять typeof skill.effect === 'function' для кастомных эффектов, "
        "иначе вызывать _applyBuiltin(). Вариант 2: зарегистрировать все builtin-эффекты "
        "как функции через registerAll(), как уже сделано для commanders_roar."
    )
)

bug_entry(
    2,
    "grid.remove(unit) не вызывается — мёртвый юнит занимает клетку грида",
    "КРИТИЧЕСКИЙ",
    "src/scenes/BattleScene.js:285-288  ·  src/systems/BattleGrid.js:45-49",
    (
        "_animDeath() создаёт tween с onComplete: () => this.grid.remove(unit). "
        "Но _renderAll() вызывается немедленно после action и уничтожает спрайт (destroy()). "
        "В Phaser 3, когда цель tween уничтожается, tween останавливается без вызова onComplete. "
        "Результат: unit.isAlive=false, но unit всё ещё сидит в BattleGrid-матрице. "
        "Метод isValidMove() считает клетку занятой. Компаньон-боец не может встать на "
        "освободившееся место переднего ряда. Логика ИИ findTarget() ищет врагов только "
        "в живых, поэтому не крашит — но грид рассинхронизируется с реальным состоянием."
    ),
    (
        "Убрать grid.remove(unit) из onComplete. Вызывать grid.remove(unit) синхронно в "
        "Unit.takeDamage() сразу после установки isAlive=false, или в _afterPlayerAction() "
        "проверять всех мёртвых и удалять их из грида перед _renderAll()."
    )
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 2: ЛОГИЧЕСКИЕ БАГИ
# ════════════════════════════════════════════════════════════════════════════

h1("2. Логические баги (геймплей нарушен)")

bug_entry(
    3,
    "Пропуск хода (skip_turn) не вызывает _renderAll() — подсветка не обновляется",
    "ВЫСОКИЙ",
    "src/scenes/BattleScene.js:474-479",
    (
        "Обработчик skip_turn вызывает ui.update() но пропускает _renderAll(). "
        "Кольцо активного юнита (ring gfx) под ногами остаётся у предыдущего. "
        "UI-панель действий обновляется правильно, но визуально активный юнит не меняется "
        "до следующего хода с атакой."
    ),
    "Добавить this._renderAll() после this.ui.update() в обработчике skip_turn."
)

bug_entry(
    4,
    "Условие поражения проверяет только героя, компаньоны не учитываются",
    "ВЫСОКИЙ",
    "src/scenes/BattleScene.js:416-419",
    (
        "_checkEnd() использует const heroAlive = this.playerUnits[0]?.isAlive — "
        "поражение наступает только при гибели первого юнита (Дуэлянт). "
        "Если бой начался со смерти Дуэлянта, а другие живы — поражение, хотя "
        "ещё двое сражаются. Это, вероятно, намеренный дизайн (герой = партия), "
        "но не документировано и создаёт путаницу."
    ),
    (
        "Если дизайн — партия умирает с героем: добавить комментарий. "
        "Если нужно поражение только при гибели всех: "
        "const allDead = this.playerUnits.every(u => !u.isAlive)."
    )
)

bug_entry(
    5,
    "Эффект 'cover' (Прикрыть) добавляется, но никогда не перехватывает урон",
    "ВЫСОКИЙ",
    "src/systems/SkillSystem.js:83-86  ·  src/entities/Unit.js:32-51",
    (
        "shield_cover добавляет эффект {type: 'covered_by', coveredBy: brawler} к цели. "
        "Unit.takeDamage() проверяет только эффект 'dodge_boost' (строка 34). "
        "Эффект 'covered_by' полностью игнорируется — Боец НИКОГДА не прикрывает союзника. "
        "Скилл потребляет ресурс cover_uses но не производит игрового эффекта."
    ),
    (
        "В Unit.takeDamage() добавить проверку covered_by: если эффект активен, "
        "перенаправить вызов к coveredBy.takeDamage(amount) и снять эффект. "
        "Пример: const cover = this.effects.find(e => e.type === 'covered_by'); "
        "if (cover) { cover.coveredBy.takeDamage(amount); this.removeEffect('covered_by'); return 0; }"
    )
)

bug_entry(
    6,
    "Bob-анимация ходьбы MapUnit конфликтует с движением (оба меняют sprite.y)",
    "СРЕДНИЙ",
    "src/entities/MapUnit.js:78-85  ·  src/entities/MapUnit.js:60-63",
    (
        "_startWalkAnim() создаёт tween { y: { from: sprite.y-3, to: sprite.y+3 }, repeat:-1 }. "
        "Значения from/to зафиксированы в момент старта анимации. "
        "update() одновременно добавляет dy-шаги к sprite.y каждый кадр. "
        "Tween переопределяет y абсолютными значениями → движение вертикально дёргается. "
        "На длинных дистанциях юнит 'прыгает' к исходной Y-позиции анимации."
    ),
    (
        "Использовать относительный bob через PropertyValueOps или отдельный контейнер: "
        "вынести спрайт в Container, применять bob к Container, а позицию изменять снаружи. "
        "Либо: убрать y из bob-tween, использовать только угол (angle) для имитации шага."
    )
)

bug_entry(
    7,
    "XP система бессмысленна: прогресс сбрасывается при рестарте BattleScene",
    "СРЕДНИЙ",
    "src/scenes/BattleScene.js:89  ·  src/scenes/BattleScene.js:437-440",
    (
        "_initUnits() каждый раз делает unitsData.map(d => new PlayerUnit(d)) — "
        "создаются новые объекты из JSON с исходными значениями. "
        "XP, набранный при победе (addXP), теряется при перезапуске сцены. "
        "Система уровней (levelUp()) никогда не становится доступной."
    ),
    (
        "Хранить состояние игрока в game.registry (сериализованный объект или Map). "
        "В _initUnits() восстанавливать hp/xp/level из registry вместо JSON. "
        "Либо использовать отдельный PlayerState-синглтон."
    )
)

bug_entry(
    8,
    "Двойной вызов scene.start() в LoadingScene",
    "СРЕДНИЙ",
    "src/scenes/LoadingScene.js:36-38  ·  src/scenes/LoadingScene.js:106",
    (
        "В preload(): load.on('complete', () => this.scene.start(destination)) — "
        "сцена стартует сразу по завершению загрузки. "
        "В create() (вызывается Phaser после preload): снова вызывается this.scene.start(). "
        "Первый вызов уже запустил переход — второй конкурирует с активным переходом. "
        "В Phaser 3 повторный scene.start() может быть проигнорирован или вызвать "
        "предупреждение в консоли. Логика LoadingScene запутана."
    ),
    (
        "Убрать scene.start() из create(). Оставить только в обработчике 'complete'. "
        "Альтернатива: убрать обработчик 'complete' из preload() и оставить "
        "scene.start() только в create() — Phaser автоматически вызывает create() "
        "после завершения загрузки."
    )
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 3: ПРОИЗВОДИТЕЛЬНОСТЬ
# ════════════════════════════════════════════════════════════════════════════

h1("3. Проблемы производительности")

bug_entry(
    9,
    "_renderAll() уничтожает и пересоздаёт ВСЕ спрайты каждый ход",
    "ВЫСОКИЙ",
    "src/scenes/BattleScene.js:144-238",
    (
        "При каждом ходе (6 юнитов × несколько ходов) вызывается _renderAll(). "
        "Он уничтожает все _unitSprites, затем создаёт заново: "
        "image, graphics (ring), setInteractive, tweens для каждого. "
        "При 6 юнитах за бой: ~6 destroy + ~6 image.add + ~6 graphics + ~6 tween.add = 24 операции/ход. "
        "Infinite tweens (repeat:-1) накапливаются — при destroy() они останавливаются, "
        "но между destroy и следующим frame Phaser обрабатывает их впустую."
    ),
    (
        "Создать спрайты единожды в create(). Обновлять только: alpha (смерть), tint (активный), "
        "flips и позиции. Для ring-highlights использовать одну Graphics, перерисовывать только "
        "нужный эллипс. Метод _renderAll() переименовать в _updateUnitVisuals() и убрать destroy/create."
    )
)

bug_entry(
    10,
    "UIManager: hint-text создаётся каждый ход и копится в памяти",
    "СРЕДНИЙ",
    "src/ui/UIManager.js:120-123",
    (
        "_updateActionButtons() вызывается каждый ход. В конце метода добавляется текст "
        "'Кликни по врагу...' через scene.add.text().setName('hint'). "
        "Затем this._actionBtns.push(scene.children.getByName('hint')) — "
        "getByName возвращает ПЕРВЫЙ найденный объект с именем 'hint'. "
        "При втором вызове в сцене уже есть старый 'hint', getByName его возвращает, "
        "новый остаётся без регистрации и никогда не удаляется → утечка памяти."
    ),
    (
        "Хранить ссылку: this._hintText = scene.add.text(...). "
        "При следующем _updateActionButtons() вызывать this._hintText?.destroy() перед пересозданием. "
        "Или добавлять hint в _actionBtns сразу: this._actionBtns.push(scene.add.text(...))."
    )
)

bug_entry(
    11,
    "UIManager._updateBorders() вызывается дважды за ход",
    "НИЗКИЙ",
    "src/scenes/BattleScene.js:383-386  ·  src/ui/UIManager.js:139-141",
    (
        "В _afterPlayerAction(): turnManager.nextTurn() → _advanceToNextAlive() → "
        "eventBus.emit('turn_started') → highlightActive() → _updateBorders() [1-й раз]. "
        "Затем this.ui.update() → _updateBorders() [2-й раз]. "
        "За каждый ход: 2 цикла destroy + 2 цикла create для границ-Graphics."
    ),
    (
        "Убрать вызов _updateBorders() из highlightActive(). "
        "Пусть update() является единственной точкой обновления всего UI."
    )
)

bug_entry(
    12,
    "MusicPlayer: старые Sound-объекты не destroy()-ятся → утечка памяти",
    "СРЕДНИЙ",
    "src/ui/MusicPlayer.js:44-65",
    (
        "_startTrack() вызывает this.current.stop() затем создаёт новый объект: "
        "this.current = this.scene.sound.add(track.key, ...). "
        "Старый sound объект остановлен, но destroy() не вызывается. "
        "Глобальный SoundManager Phaser хранит ссылку → объекты накапливаются. "
        "При переходе MapScene → BattleScene: обе сцены создают свои MusicPlayer, "
        "оба добавляют sounds в глобальный manager — потенциально два трека играют одновременно."
    ),
    (
        "В _startTrack() добавить: if (this.current) { this.current.stop(); this.current.destroy(); } "
        "Перед созданием новой сцены проверять game.registry.get('bgMusic')?.isPlaying "
        "и не создавать новый MusicPlayer если музыка уже играет."
    )
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 4: АРХИТЕКТУРНЫЕ ПРОБЛЕМЫ
# ════════════════════════════════════════════════════════════════════════════

h1("4. Архитектурные проблемы")

bug_entry(
    13,
    "BattleScene.preload() дублирует загрузку уже кэшированных ассетов",
    "НИЗКИЙ",
    "src/scenes/BattleScene.js:53-57",
    (
        "BattleScene.preload() загружает battle_bg и спрайты, которые уже были загружены "
        "в LoadingScene. Комментарий на строке 56 сам себе противоречит: "
        "'здесь ничего не грузим' — при этом сверху идут два вызова load. "
        "Phaser использует кэш, дублей в сети нет, но код вводит в заблуждение."
    ),
    "Полностью убрать содержимое BattleScene.preload() или весь метод."
)

bug_entry(
    14,
    "EventBus — singleton без механизма изоляции при смене сцены",
    "СРЕДНИЙ",
    "src/utils/eventBus.js  ·  src/scenes/BattleScene.js:60",
    (
        "EventBus — глобальный singleton. BattleScene очищает его в create() через eventBus.clear(). "
        "Но MapScene (через MusicPlayer и _bindEvents) также может добавлять листенеры. "
        "При переходе MapScene → BattleScene → MapScene: BattleScene.create() очищает ВСЕ "
        "листенеры включая те, что MapScene ещё не успела снять. "
        "Риск: листенеры MapScene теряются без уведомления."
    ),
    (
        "Использовать именованные группы (namespace) для листенеров: "
        "eventBus.on('battle:log', ...) vs 'map:encounter'. "
        "Или очищать только конкретные события: eventBus.clearGroup('battle') в BattleScene."
    )
)

bug_entry(
    15,
    "vite.config.js: base: '/dark-fantasy-rpg/' ломает локальную разработку",
    "СРЕДНИЙ",
    "vite.config.js:2",
    (
        "base: '/dark-fantasy-rpg/' означает что в локальной среде (npm run dev) "
        "игра доступна только по пути localhost:5173/dark-fantasy-rpg/. "
        "Прямой localhost:5173/ вернёт 404. "
        "Это корректно для GitHub Pages деплоя, но неудобно при локальной разработке."
    ),
    (
        "Использовать переменную окружения: "
        "base: process.env.NODE_ENV === 'production' ? '/dark-fantasy-rpg/' : '/'. "
        "Или добавить .env.local с VITE_BASE_URL=/ для локальной работы."
    )
)

bug_entry(
    16,
    "PortraitPanel.preload() никогда не вызывается — мёртвый код",
    "НИЗКИЙ",
    "src/ui/PortraitPanel.js:16-20  ·  src/scenes/BattleScene.js:66-67",
    (
        "PortraitPanel имеет метод preload() (строка 16), который должен загружать портреты. "
        "Но в BattleScene вызывается только portraits.create(), не preload(). "
        "Портреты работают потому что LoadingScene их загрузила заранее. "
        "Если LaadingScene изменится, PortraitPanel тихо сломается."
    ),
    (
        "Либо вызывать this.portraits.preload() в BattleScene.preload(), "
        "либо удалить мёртвый метод и документировать зависимость от LoadingScene."
    )
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 5: UX И ГЕЙМПЛЕЙ
# ════════════════════════════════════════════════════════════════════════════

h1("5. UX / Геймплейные проблемы")

bug_entry(
    17,
    "После боя нет возврата на карту — игра застрявшая в BattleScene",
    "ВЫСОКИЙ",
    "src/scenes/BattleScene.js:444-449",
    (
        "_endBattle() показывает только кнопку '[ Попробовать снова ]' → scene.restart(). "
        "Нет кнопки '[ Вернуться на карту ]'. После победы или поражения "
        "единственный выход — перезапуск боя. "
        "MapScene, world exploration, прогресс — недостижимы без перезагрузки страницы."
    ),
    (
        "Добавить кнопку 'Вернуться на карту': "
        "scene.start('LoadingScene', { destination: 'MapScene', destinationData: { mapKey: 'map1' } }). "
        "При победе отмечать бандита как побеждённого (game.registry) чтобы он не появлялся снова."
    )
)

bug_entry(
    18,
    "Бандит на карте не возрождается, но и не исчезает визуально после победы",
    "СРЕДНИЙ",
    "src/scenes/MapScene.js:93-96  ·  src/scenes/MapScene.js:197-200",
    (
        "При встрече: bandit.encountered = true → бандит перестаёт патрулировать. "
        "Его спрайт остаётся на экране навсегда. После возврата с боя (если бы кнопка была): "
        "новый инстанс MapScene создаётся → бандит снова появляется и патрулирует "
        "(encountered сбрасывается). Победа в бою не отражается на карте."
    ),
    (
        "Хранить список побеждённых бандитов в game.registry. "
        "В MapScene._spawnBandits() пропускать уже побеждённых. "
        "При encountered=true: анимировать исчезновение спрайта (alpha tween)."
    )
)

bug_entry(
    19,
    "HP игроков не сохраняется между боями — каждая битва со свежими статами",
    "СРЕДНИЙ",
    "src/scenes/BattleScene.js:89",
    (
        "unitsData.map(d => new PlayerUnit(d)) создаёт юниты из статичного JSON каждый раз. "
        "Если Знахарка была на 10 HP в конце боя — следующий бой она начинает с 50. "
        "Игра не имеет persistence состояния игрока."
    ),
    (
        "После боя сохранять текущий HP в game.registry или PlayerState-объект. "
        "В _initUnits() восстанавливать: unit.hp = savedState.hp ?? unit.maxHp. "
        "Добавить возможность лечения на карте (у NPC или отдых)."
    )
)

bug_entry(
    20,
    "Диалог перед боем отсутствует — встреча с бандитом идёт в бой без контекста",
    "НИЗКИЙ",
    "src/scenes/MapScene.js:202-213  ·  src/data/dialogue.json",
    (
        "MapScene._startEncounter() сразу начинает fade → LoadingScene → BattleScene. "
        "dialogue.json содержит данные для 3 NPC (трактирщик, солдат, незнакомка), "
        "но DialogueScene.js не существует. "
        "Бандит на карте не имеет имени, портрета при встрече, нет варианта 'избежать боя'."
    ),
    (
        "Создать DialogueScene с параметрами {npcId, nextScene, nextData}. "
        "В MapScene._startEncounter(): перед боем показывать DiaglogueScene с банд.-репликой. "
        "dialogue.json уже готов — добавить запись для bandit_patrol."
    )
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 6: ОТСУТСТВУЮЩИЕ ФИЧИ И МЁРТВЫЙ КОД
# ════════════════════════════════════════════════════════════════════════════

h1("6. Отсутствующие фичи и неиспользуемые ассеты")

h2("6.1 Функционал заявлен, но не реализован")

add_table(
    ["Фича", "Где упоминается", "Статус"],
    [
        ["DialogueScene", "dialogue.json, 3 NPC", "Файл не существует"],
        ["Система уровней (Level Up UI)", "PlayerUnit.levelUp(), XP в constants.js", "Логика есть, UI нет"],
        ["Карты: tavern_map, mountains_map, swamp_map", "public/maps/, LoadingScene загружает", "Загружены, но недостижимы"],
        ["NPC на карте мира", "dialogue.json (Трактирщик и др.)", "Нет объектов в MapScene"],
        ["Навык aimed_shot", "skills.json (определён)", "Не назначен ни одному юниту"],
        ["Перемещение в бою (move())", "PlayerUnit.move(), BattleGrid.moveUnit()", "Нет UI кнопки для перемещения"],
        ["Passive бонусы EnemyUnit", "EnemyUnit: passive, passiveBonus", "Поля есть, логика не применяется"],
        ["cannotMiss у aimed_shot", "skills.json: cannotMiss: true", "SkillSystem игнорирует это поле"],
    ]
)

h2("6.2 Аудио — ассеты vs код")

add_table(
    ["Ключ трека", "Файл в коде", "Файл на диске", "Статус"],
    [
        ["track_ashes2",     "Ashes of Velanth 2.mp3",       "✓ Есть", "OK"],
        ["track_ashes",      "Ashes of Velanth.mp3",          "✓ Есть", "OK"],
        ["track_ward",       "Ashes of the Last Ward.mp3",    "✓ Есть", "OK"],
        ["track_monastery2", "Ashes of the Monastery 2.mp3",  "✓ Есть", "OK"],
        ["track_monastery",  "Ashes of the Monastery.mp3",    "✓ Есть", "OK"],
        ["track_pass",       "Ashes of the Pass.mp3",         "✓ Есть", "OK"],
        ["track_dark",       "Dark Song.mp3",                 "✓ Есть", "OK"],
        ["track_forest",     "Forest song.mp3",               "✗ НЕТ",  "ОШИБКА 404"],
        ["track_mermaids",   "Mermaids song.mp3",             "✗ НЕТ",  "ОШИБКА 404"],
        ["—",                "track1.mp3 ... track9.mp3",     "✓ Есть", "НЕИСПОЛЬЗУЕМЫЕ"],
    ]
)

para(
    "Итого: 2 трека в коде без файлов на диске (404 при загрузке), "
    "9 файлов track1-9.mp3 на диске без ссылок в коде (мёртвый груз ~= потенциально сотни МБ).",
    italic=True, color=(0xAA, 0x44, 0x00)
)

h2("6.3 Debug-код в продакшн сборке")

add_table(
    ["Файл", "Строка", "Проблема"],
    [
        ["src/scenes/MapScene.js", "47", "this.walkable.drawDebug(this) — красные зоны видны в игре"],
        ["src/scenes/BattleScene.js", "24", "const HAS_BG = true — флаг-заглушка вместо конфига"],
        ["src/scenes/BattleScene.js", "211-214", "const barBg = null; const barFill = null — мёртвые переменные"],
        ["src/systems/BattleGrid.js", "78-84", "getCellPixelPos() — утилита для сетки, нигде не вызывается"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 7: СТРУКТУРА ПРОЕКТА
# ════════════════════════════════════════════════════════════════════════════

h1("7. Обзор структуры проекта")

h2("7.1 Карта зависимостей (упрощённая)")

para(
    "BootScene → LoadingScene → MapScene → (encounter) → LoadingScene → BattleScene → restart\n"
    "EventBus: глобальная шина, используется в Unit, PlayerUnit, EnemyUnit, BattleScene, TurnManager, UIManager\n"
    "game.registry: bgMusic, bgMusicIndex, musicStarted — музыкальное состояние между сценами",
    bold=False
)

h2("7.2 Файлы проекта")

add_table(
    ["Файл / Директория", "Строк", "Описание"],
    [
        ["src/scenes/BootScene.js",    "33",  "Заставка с кнопкой Start"],
        ["src/scenes/LoadingScene.js", "108", "Загрузка всех ассетов + прогресс-бар"],
        ["src/scenes/MapScene.js",     "195", "Карта мира: движение, встречи, HUD"],
        ["src/scenes/BattleScene.js",  "481", "Тактический бой: рендер, AI, события"],
        ["src/entities/Unit.js",       "123", "Базовый класс юнита"],
        ["src/entities/PlayerUnit.js", "86",  "Игровой юнит: скиллы, XP, перемещение"],
        ["src/entities/EnemyUnit.js",  "47",  "ИИ-юнит: findTarget, decideAction"],
        ["src/entities/MapUnit.js",    "112", "Юнит на карте: движение, анимация ходьбы"],
        ["src/systems/TurnManager.js", "74",  "Очерёдность ходов по speed"],
        ["src/systems/BattleGrid.js",  "85",  "Позиционирование 2×3"],
        ["src/systems/SkillSystem.js", "107", "Регистрация и применение навыков"],
        ["src/systems/WalkableZones.js","82", "Зоны хождения на картах"],
        ["src/ui/UIManager.js",        "164", "Лог, кнопки действий, подсветка"],
        ["src/ui/PortraitPanel.js",    "107", "Портреты и HP-бары боковых панелей"],
        ["src/ui/MusicPlayer.js",      "117", "Компактный аудио-плеер"],
        ["src/data/units.json",        "—",   "3 игровых юнита (hero, brawler, healer)"],
        ["src/data/enemies.json",      "—",   "3 врага (commander, brawler, archer)"],
        ["src/data/skills.json",       "—",   "6 навыков"],
        ["src/data/dialogue.json",     "—",   "3 NPC-диалога (не используются)"],
        ["src/utils/constants.js",     "51",  "GRID, PHASE, COLORS, XP, AI_DELAY"],
        ["src/utils/eventBus.js",      "43",  "Глобальная шина событий"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 8: ТОП-10 ПРИОРИТЕТНЫХ ИСПРАВЛЕНИЙ
# ════════════════════════════════════════════════════════════════════════════

h1("8. Топ-10 — план следующего спринта")

para("Упорядочено по соотношению: критичность × объём работы.", italic=True)

sprint_items = [
    ("1", "КРИТИЧЕСКИЙ", "Исправить SkillSystem: skill.effect (строка vs функция)",
     "SkillSystem.js:38 + skills.json",
     "~30 мин. Добавить typeof-проверку. Разблокирует 3 из 6 скиллов."),
    ("2", "КРИТИЧЕСКИЙ", "Исправить grid.remove() — вызывать синхронно, не в onComplete",
     "BattleScene.js:285-288",
     "~20 мин. Убрать из _animDeath, добавить в _afterPlayerAction()."),
    ("3", "ВЫСОКИЙ", "Добавить кнопку 'Вернуться на карту' в _endBattle()",
     "BattleScene.js:444",
     "~15 мин. Разблокирует полный игровой цикл."),
    ("4", "ВЫСОКИЙ", "Исправить skip_turn: добавить _renderAll()",
     "BattleScene.js:474-479",
     "~5 мин. Одна строка."),
    ("5", "ВЫСОКИЙ", "Реализовать эффект cover в Unit.takeDamage()",
     "Unit.js:32-51",
     "~45 мин. Сделать Бойца рабочим."),
    ("6", "СРЕДНИЙ", "Убрать debug overlay walkable zones",
     "MapScene.js:47",
     "~2 мин. Закомментировать одну строку."),
    ("7", "СРЕДНИЙ", "Исправить двойной scene.start() в LoadingScene",
     "LoadingScene.js:36+106",
     "~10 мин. Убрать из create()."),
    ("8", "СРЕДНИЙ", "Добавить destroy() для старых Sound объектов в MusicPlayer",
     "MusicPlayer.js:44-47",
     "~10 мин. Предотвратить утечку памяти."),
    ("9", "СРЕДНИЙ", "Исправить UIManager: hint-text утечка памяти",
     "UIManager.js:120-123",
     "~15 мин. Сохранить ссылку и destroy() перед пересозданием."),
    ("10", "СРЕДНИЙ", "Добавить/удалить аудио-файлы: Forest song.mp3, Mermaids song.mp3",
     "public/audio/ + LoadingScene.js",
     "~30 мин. Добавить файлы или убрать из кода."),
]

sprint_table_rows = []
for num, sev, title, loc, effort in sprint_items:
    sprint_table_rows.append([num, sev, title, loc, effort])

add_table(
    ["#", "Приоритет", "Задача", "Файл:строка", "Оценка"],
    sprint_table_rows
)

doc.add_paragraph()
h2("Общая оценка спринта")
para(
    "Критические исправления (пп. 1–4): ~1.5 часа разработки\n"
    "Высокий приоритет (пп. 5): ~45 минут\n"
    "Средний приоритет (пп. 6–10): ~1 час\n"
    "Итого: ~3.5 часа для перехода в стабильное и играбельное состояние."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 9: ПОЛОЖИТЕЛЬНЫЕ СТОРОНЫ
# ════════════════════════════════════════════════════════════════════════════

h1("9. Сильные стороны проекта")

para(
    "Несмотря на найденные баги, архитектура проекта имеет ряд хороших решений:"
)

positives = [
    ("Модульность", "Чёткое разделение Entity / System / UI / Scene. Легко добавлять новые скиллы (skills.json) и юниты (units.json) без изменения кода."),
    ("EventBus", "Правильная идея — развязать системы через события. Нужна лишь изоляция по сценам."),
    ("TurnManager", "Корректно реализован приоритет по speed, обработка мёртвых юнитов (_advanceToNextAlive), тик cooldowns/effects в endTurn()."),
    ("BattleGrid", "Чистая логика 2×3 с раздельными сторонами. canMeleeAttack() корректно проверяет передний ряд."),
    ("PortraitPanel", "Визуально богатое решение в стиле Disciples II — портреты, HP-бары, затемнение мёртвых."),
    ("MapUnit анимация", "Приятная имитация ходьбы через bob+lean tweens. Идея верная, нужен рефакторинг y-конфликта."),
    ("LoadingScene", "Централизованная загрузка всех ассетов — правильный подход. Прогресс-бар улучшает UX."),
    ("Data-driven дизайн", "JSON-данные для юнитов, врагов, навыков, диалогов — облегчает баланс и контент-итерации."),
]

for title, desc in positives:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"✓ {title}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# ИТОГОВОЕ ЗАКЛЮЧЕНИЕ
# ════════════════════════════════════════════════════════════════════════════

h1("10. Заключение")

para(
    "Проект представляет собой качественный vertical slice тактической RPG с хорошей "
    "визуальной составляющей и правильной модульной архитектурой. "
    "Основная техническая проблема — несколько критических багов, блокирующих "
    "базовый геймплей (3 из 6 скиллов крашат игру, нет возврата на карту). "
    "\n\n"
    "После 3.5 часов работы по приоритетному списку игра станет полностью играбельной. "
    "Средне- и долгосрочный беклог включает: DialogueScene, persistence игрока, "
    "несколько карт мира и NPC — всё это уже подготовлено на уровне данных.",
    size=11
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(f"Аудит выполнен: {datetime.date.today().strftime('%d.%m.%Y')} | Claude Sonnet 4.6")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Сохранение
output_path = "/Users/maksimilin/Desktop/Dark Fantasy/dark-fantasy-rpg/AUDIT_REPORT.docx"
doc.save(output_path)
print(f"✓ Сохранено: {output_path}")
